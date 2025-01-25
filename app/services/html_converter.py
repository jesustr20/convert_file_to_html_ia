from docx import Document
from PyPDF2 import PdfReader
import pandas as pd
import pdfplumber
import os

from app.services.file_handler import save_file

def convert_file_to_html(file):
    """
    Convierte un archivo DOCX, PDF o Excel a HTML.
    """

    file_location = save_file(file)

    file_extension = file.filename.split('.')[-1].lower()

    if file_extension == "docx":
        file_to_html = convert_docx_to_html(file_location)
    elif file_extension == "pdf":
        file_to_html = convert_pdf_to_html(file_location)
    elif file_extension == "xlsx":
        file_to_html =  convert_excel_to_html(file_location)
    else:
        raise ValueError("Unsupported file type. Only DOCX, PDF, and XLSX are allowed.")
    
    os.remove(file_location)

    return file_to_html

def convert_docx_to_html(file_path: str) -> str:

    doc = Document(file_path)
    texto = []

    for paragraph in doc.paragraphs:
        parrafo = paragraph.text
        if paragraph.style.font.bold:
             parrafo = "[Negrita]" + parrafo
        texto.append(parrafo)
    
    for tabla in doc.tables:        
        for fila in tabla.rows:
            for celda in fila.cells:
                texto.append("[Tabla]" + celda.text)

    return "\n".join(texto)

#En proceso de construccion
def convert_pdf_to_html(file_path):
    with pdfplumber.open(file_path) as pdf:  
        html = "<html><body>"
        
        for page in pdf.pages:
            html += f"<div>{page.extract_text()}</div>"
    
        html += "</body></html>"
    return html

def convert_excel_to_html(file_path):
    df = pd.read_excel(file_path)
    html = "<html><body><table border='1'>"

    html += "<tr>"
    for column in df.columns:
        html += f"<th>{column}></th>"
    html += "</tr>"

    for index, row in df.iterrows():
        html += "<tr>"
        for value in row:
            html += f"<td>{value}</td>"
        html += "</tr>"

    html += "</table></body></html>"
    return html